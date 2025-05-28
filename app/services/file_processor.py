from fastapi import UploadFile, HTTPException
import os
import uuid
from typing import Tuple, Dict, Any
import aiofiles
import magic
import PyPDF2
import io
import tempfile
from docx import Document

class FileProcessor:
    
    async def process_file(self, file: UploadFile) -> Dict[str, Any]:
        
        content = await file.read(1024)  
        file_type = magic.from_buffer(content, mime=True)
        await file.seek(0) 
        
        allowed_types = ['application/pdf', 
                        'application/msword', 
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document']
        
        print(f"File type: {file_type}")
        if file_type not in allowed_types:
                if file.filename:
                    filename_lower = file.filename.lower()
                    if filename_lower.endswith('.pdf'):
                        detected_type = 'application/pdf'
                    elif filename_lower.endswith('.docx'):
                        detected_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    elif filename_lower.endswith('.doc'):
                        detected_type = 'application/msword'
        
        if detected_type not in allowed_types:
            raise HTTPException(status_code=400, detail="Not supported file type, only PDF and Word documents are supported for now")
        
        
        pages = await self._count_pages(file, detected_type)
        
        return {
            "filename": None, 
            "original_name": file.filename,
            "pages": pages,
            "content_type": file_type
        }
    
    async def _count_pages(self, file: UploadFile, file_type: str) -> int:
        try:
            await file.seek(0)
            content = await file.read()
            await file.seek(0)
            
            print(f"File name: {file.filename}")
            print(f"Detected MIME type: {file_type}")
            print(f"Content type from upload: {file.content_type}")
            
            if file_type == 'application/pdf':
                pdf_file = io.BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                return len(pdf_reader.pages)
            elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return await self._estimate_word_pages(content, file_type)
            elif file_type == 'application/msword':
                 file_size_kb = len(content) / 1024
                 return max(1, int(file_size_kb / 75))
            else:
                raise HTTPException(status_code=400, detail="Not supported file type, only PDF and Word documents are supported for now")
        except Exception as e:
            print(f"Error reading document: {str(e)}")
            raise HTTPException(status_code=500, detail="Error reading document")
    

    # async def _count_word_pages_via_pdf(self, content: bytes, file_type: str) -> int:
    #     temp_dir = None
    #     try:
    #         temp_dir = tempfile.mkdtemp()
            
    #         if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
    #             ext = '.docx'
    #         else:
    #             ext = '.doc'
            
    #         word_path = os.path.join(temp_dir, f"document{ext}")
    #         with open(word_path, 'wb') as f:
    #             f.write(content)
            
    #         pdf_path = os.path.join(temp_dir, "document.pdf")
            
    #         # 使用 pandoc 转换
    #         cmd = [
    #             'pandoc',
    #             word_path,
    #             '-o', pdf_path
    #         ]
            
    #         process = await asyncio.create_subprocess_exec(
    #             *cmd,
    #             stdout=asyncio.subprocess.PIPE,
    #             stderr=asyncio.subprocess.PIPE
    #         )
            
    #         stdout, stderr = await process.communicate()
            
    #         if process.returncode != 0:
    #             print(f"Pandoc conversion failed: {stderr.decode()}")
    #             return await self._estimate_word_pages(content, file_type)
            
    #         if os.path.exists(pdf_path):
    #             with open(pdf_path, 'rb') as pdf_file:
    #                 pdf_reader = PyPDF2.PdfReader(pdf_file)
    #                 return len(pdf_reader.pages)
    #         else:
    #             return await self._estimate_word_pages(content, file_type)
                
    #     except Exception as e:
    #         print(f"Error converting Word to PDF: {str(e)}")
    #         return await self._estimate_word_pages(content, file_type)
    #     finally:
    #         if temp_dir and os.path.exists(temp_dir):
    #             import shutil
    #             shutil.rmtree(temp_dir, ignore_errors=True)


    async def _estimate_word_pages(self, content: bytes, file_type: str) -> int:
        try:
            if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    
                    doc = Document(temp_file.name)
                    
                    
                    total_lines = 0
                    total_words = 0
                    
                    # count paragraphs
                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if text:  # non-empty paragraph
                            words = len(text.split())
                            total_words += words
                            
                            # estimate lines
                            lines_in_paragraph = max(1, words // 13)
                            total_lines += lines_in_paragraph
                            
                            # paragraph spacing
                            total_lines += 0.5
                    
                    # count table lines
                    for table in doc.tables:
                        rows = len(table.rows)
                        cols = len(table.columns) if table.rows else 0
                        
                        # table lines + empty lines before and after table
                        table_lines = rows * 1.2 + 2
                        total_lines += table_lines
                    
                    # count images lines
                    image_count = 0
                    for paragraph in doc.paragraphs:
                        for run in paragraph.runs:
                            if hasattr(run.element, 'xpath'):
                                images = run.element.xpath('.//a:blip')
                                image_count += len(images)
                    
                    # count images lines
                    total_lines += image_count * 4
                    
                    # A4 page has about 54 lines (12pt font, single line spacing)
                    lines_per_page = 54
                    estimated_pages = max(1, total_lines / lines_per_page)
                    
                    # margins and formatting factor
                    adjustment_factor = 1.1
                    final_pages = max(1, int(estimated_pages * adjustment_factor))
                    
                    os.unlink(temp_file.name)
                    return final_pages
                    
            else:
                # .doc file
                file_size_kb = len(content) / 1024
                
                # use different estimation ratios based on file size
                if file_size_kb < 50:  # small file
                    pages_per_kb = 1/30 
                elif file_size_kb < 200:  # medium file
                    pages_per_kb = 1/50
                else:  # large file
                    pages_per_kb = 1/75
                
                estimated_pages = max(1, int(file_size_kb * pages_per_kb))
                return estimated_pages
                
        except Exception as e:
            print(f"Error in page estimation: {str(e)}")
            # downgrade estimation
            if file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                try:
                    # try basic docx parsing
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                        temp_file.write(content)
                        temp_file.flush()
                        
                        doc = Document(temp_file.name)
                        
                        total_text = ""
                        for paragraph in doc.paragraphs:
                            total_text += paragraph.text + "\n"
                        
                        chars_per_page = 2500
                        estimated_pages = max(1, len(total_text) // chars_per_page)
                        
                        table_pages = len(doc.tables) * 0.3
                        
                        os.unlink(temp_file.name)
                        return max(1, int(estimated_pages + table_pages))
                except:
                    # downgrade to file size calculation
                    return max(1, len(content) // 40960)
            else:
                # .doc file downgrade
                file_size_kb = len(content) / 1024
                return max(1, int(file_size_kb / 75))

    async def save_file(self, file: UploadFile, upload_folder: str) -> str:
        os.makedirs(upload_folder, exist_ok=True)
        
        file_extension = os.path.splitext(file.filename)[1] if file.filename else ""
        file_id = uuid.uuid4()
        file_path = os.path.join(upload_folder, f"{file_id}{file_extension}")
        
        await file.seek(0)
        async with aiofiles.open(file_path, 'wb') as out_file:
            while content := await file.read(1024 * 1024):  
                await out_file.write(content)
        
        return file_id, file_path

file_processor = FileProcessor()

def get_file_processor():
    return file_processor